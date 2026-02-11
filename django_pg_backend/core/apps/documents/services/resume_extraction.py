"""
Resume Extraction Engine
Handles extraction of raw text from PDF and DOCX files.
"""

import os
import re
from typing import Optional, Tuple
import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument


class ResumeExtractionEngine:
    """
    Extracts raw text from resume files (PDF/DOCX).
    """
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.doc']
    
    def extract_text(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Extract text from a resume file based on its extension.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        if not os.path.exists(file_path):
            return None, f"File not found: {file_path}"
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        else:
            return None, f"Unsupported file format: {file_ext}"
    
    def _extract_from_pdf(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Extract text from PDF using PyMuPDF (fitz) as primary, pdfplumber as fallback.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            # Try PyMuPDF first (faster)
            text = self._extract_with_pymupdf(file_path)
            if text and len(text.strip()) > 50:
                return self._clean_text(text), None
            
            # Fallback to pdfplumber (better for some PDFs)
            text = self._extract_with_pdfplumber(file_path)
            if text and len(text.strip()) > 50:
                return self._clean_text(text), None
            
            return None, "Could not extract sufficient text from PDF"
            
        except Exception as e:
            return None, f"PDF extraction error: {str(e)}"
    
    def _extract_with_pymupdf(self, file_path: str) -> str:
        """
        Extract text using PyMuPDF (fitz).
        """
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception:
            pass
        return text
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """
        Extract text using pdfplumber.
        """
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            pass
        return text
    
    def _extract_from_docx(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if len(text.strip()) > 50:
                return self._clean_text(text), None
            else:
                return None, "Could not extract sufficient text from DOCX"
                
        except Exception as e:
            return None, f"DOCX extraction error: {str(e)}"
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and special characters.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Remove bullet point artifacts
        text = re.sub(r'[•●○■▪▫]', '', text)
        
        # Clean up common PDF extraction artifacts
        text = re.sub(r'\f', '\n\n', text)  # Form feed to newlines
        text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)  # Fix hyphenated words
        
        # Remove multiple consecutive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def is_supported_file(self, file_path: str) -> bool:
        """
        Check if the file format is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if supported, False otherwise
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        return file_ext in self.supported_extensions
